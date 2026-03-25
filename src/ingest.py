import os
import pdfplumber
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from supabase import create_client
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument

# ── Load environment variables from .env ──────────────────────────────────────
load_dotenv()

openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
supabase = create_client(os.getenv("SUPABASE_URL"), os.getenv("SUPABASE_KEY"))

# ── Configuration ─────────────────────────────────────────────────────────────
DOCS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
DOCS_PATH = Path(DOCS_DIR)
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
EMBEDDING_MODEL = "text-embedding-3-small"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
PARAGRAPHS_PER_PAGE = 30   # virtual page size for Word docs
CHARS_PER_PAGE = 2000      # virtual page size for plain text

# ── Prompt injection patterns to block ────────────────────────────────────────
INJECTION_PATTERNS = [
    "ignore previous instructions",
    "ignore all previous",
    "disregard your instructions",
    "forget everything above",
    "new instructions:",
    "you are now",
    "act as if you",
    "system prompt:",
    "override your",
    "jailbreak",
]

def check_prompt_injection(pages: list[dict]) -> bool:
    """Return True if the extracted text contains prompt injection patterns."""
    all_text = " ".join(p["text"] for p in pages).lower()
    return any(pattern in all_text for pattern in INJECTION_PATTERNS)


def check_relevance(pages: list[dict]) -> bool:
    """Return True if the document is relevant to the construction knowledge base."""
    # Sample up to the first 1500 chars to keep the check fast and cheap
    sample = " ".join(p["text"] for p in pages)[:1500]

    response = openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a document relevance classifier for a construction industry knowledge base "
                    "that covers commercial roofing, glazing, curtain walls, building envelope systems, "
                    "building codes, site safety, and related technical specifications.\n\n"
                    "Respond with ONLY 'yes' if the document is relevant to any of those topics, "
                    "or 'no' if it is not (e.g. invoices, personal documents, unrelated business docs)."
                )
            },
            {"role": "user", "content": f"Document sample:\n{sample}"}
        ],
        temperature=0,
        max_tokens=3
    )

    answer = response.choices[0].message.content.strip().lower()
    return answer == "yes"

# ── Step 1: Extract text from a PDF ───────────────────────────────────────────
def extract_text_from_pdf(pdf_path: Path) -> list[dict]:
    """Extract text page by page from a PDF. Returns list of {page, text} dicts."""
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                pages.append({
                    "page": page_num,
                    "text": text.strip()
                })
    return pages


# ── Extract text from a Word document ─────────────────────────────────────────
def extract_text_from_docx(docx_path: Path) -> list[dict]:
    """Extract text from a .docx file. Groups paragraphs into virtual pages."""
    doc = DocxDocument(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    pages = []
    for i in range(0, len(paragraphs), PARAGRAPHS_PER_PAGE):
        chunk = "\n".join(paragraphs[i: i + PARAGRAPHS_PER_PAGE])
        if chunk:
            pages.append({"page": i // PARAGRAPHS_PER_PAGE + 1, "text": chunk})
    return pages


# ── Extract text from a plain text file ───────────────────────────────────────
def extract_text_from_txt(txt_path: Path) -> list[dict]:
    """Extract text from a .txt file. Splits into virtual pages by character count."""
    text = txt_path.read_text(encoding="utf-8", errors="ignore")
    pages = []
    for i in range(0, len(text), CHARS_PER_PAGE):
        chunk = text[i: i + CHARS_PER_PAGE].strip()
        if chunk:
            pages.append({"page": i // CHARS_PER_PAGE + 1, "text": chunk})
    return pages


# ── Dispatch extractor by file type ───────────────────────────────────────────
def extract_text(file_path: Path) -> list[dict]:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    return []


# ── Step 2: Split pages into chunks ───────────────────────────────────────────
def chunk_pages(pages: list[dict], filename: str) -> list[dict]:
    """Split page text into overlapping chunks. Preserves metadata."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " "]
    )

    chunks = []
    for page in pages:
        splits = splitter.split_text(page["text"])
        for i, split in enumerate(splits):
            chunks.append({
                "content": split,
                "metadata": {
                    "filename": filename,
                    "page": page["page"],
                    "chunk_index": i
                }
            })
    return chunks


# ── Step 3: Embed a batch of chunks ───────────────────────────────────────────
def embed_chunks(chunks: list[dict]) -> list[dict]:
    """Call OpenAI embeddings API and attach embedding to each chunk."""
    texts = [chunk["content"] for chunk in chunks]

    response = openai_client.embeddings.create(
        model=EMBEDDING_MODEL,
        input=texts
    )

    for i, embedding_obj in enumerate(response.data):
        chunks[i]["embedding"] = embedding_obj.embedding

    return chunks


# ── Step 4: Store chunks in Supabase ──────────────────────────────────────────
BATCH_SIZE = 50

def store_chunks(chunks: list[dict], tenant_id: str = "default",
                  version: int = 1, is_latest: bool = True, original_filename: str = ""):
    """Insert chunks into the documents table in Supabase in batches."""
    rows = [
        {
            "content": chunk["content"],
            "metadata": chunk["metadata"],
            "embedding": chunk["embedding"],
            "tenant_id": tenant_id,
            "version": version,
            "is_latest": is_latest,
            "original_filename": original_filename or chunk["metadata"].get("filename", ""),
        }
        for chunk in chunks
    ]
    for i in range(0, len(rows), BATCH_SIZE):
        batch = rows[i: i + BATCH_SIZE]
        supabase.table("documents").insert(batch).execute()
    print(f"  Stored {len(rows)} chunks (version={version})")


# ── Step 5: Register document in registry ─────────────────────────────────────
def register_document(filename: str, chunk_count: int, tenant_id: str = "default",
                       version: int = 1, is_latest: bool = True, original_filename: str = ""):
    """Record the document in document_registry with upload timestamp."""
    supabase.table("document_registry").upsert({
        "filename": filename,
        "chunk_count": chunk_count,
        "tenant_id": tenant_id,
        "version": version,
        "is_latest": is_latest,
        "original_filename": original_filename or filename,
    }, on_conflict="filename,tenant_id").execute()


# ── Main ingestion loop ────────────────────────────────────────────────────────
def ingest_all():
    all_files = [
        f for ext in SUPPORTED_EXTENSIONS
        for f in DOCS_PATH.glob(f"*{ext}")
    ]

    if not all_files:
        print("No supported files found in /docs folder.")
        return

    print(f"Found {len(all_files)} files\n")

    for file_path in all_files:
        existing = supabase.table("documents") \
            .select("id") \
            .eq("metadata->>filename", file_path.name) \
            .limit(1) \
            .execute()
        if existing.data:
            print(f"Skipping (already ingested): {file_path.name}")
            continue

        print(f"Processing: {file_path.name}")
        pages = extract_text(file_path)
        print(f"  Extracted {len(pages)} pages")

        chunks = chunk_pages(pages, file_path.name)
        print(f"  Created {len(chunks)} chunks")

        chunks = embed_chunks(chunks)
        print(f"  Embedded {len(chunks)} chunks")

        store_chunks(chunks, tenant_id="default", version=1, is_latest=True, original_filename=file_path.name)
        register_document(file_path.name, len(chunks), tenant_id="default",
                          version=1, is_latest=True, original_filename=file_path.name)
        print(f"  Done: {file_path.name}\n")

    print("Ingestion complete.")


def get_existing_versions(filename: str, tenant_id: str = "default") -> list[dict]:
    """Return all existing version records for a filename from document_registry."""
    result = supabase.table("document_registry") \
        .select("filename, version, is_latest, uploaded_at, chunk_count") \
        .eq("original_filename", filename) \
        .eq("tenant_id", tenant_id) \
        .order("version", desc=False) \
        .execute()
    return result.data


def get_document_version_history(filename: str, tenant_id: str = "default") -> list[dict]:
    """Return all versions of a document with their version numbers and upload dates."""
    return get_existing_versions(filename, tenant_id)


def ingest_single_file(file_path: str, filename: str, tenant_id: str = "default",
                        force_new_version: bool = False):
    """Ingest a single file — used for runtime uploads from the UI.
    If force_new_version=True and a document with the same filename exists,
    mark old versions as not latest and ingest as a new version."""
    ext = Path(file_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return False, f"Unsupported file type '{ext}'. Accepted: {', '.join(SUPPORTED_EXTENSIONS)}"

    print(f"Ingesting: {filename} (tenant: {tenant_id}, force_new_version={force_new_version})")

    existing_versions = get_existing_versions(filename, tenant_id)

    if existing_versions and not force_new_version:
        return False, f"{filename} has already been ingested. No duplicates added."

    # Determine version number
    if existing_versions:
        max_version = max(v["version"] for v in existing_versions)
        new_version = max_version + 1
        # Mark all existing chunks and registry entries as not latest
        supabase.table("documents") \
            .update({"is_latest": False}) \
            .eq("original_filename", filename) \
            .eq("tenant_id", tenant_id) \
            .execute()
        supabase.table("document_registry") \
            .update({"is_latest": False}) \
            .eq("original_filename", filename) \
            .eq("tenant_id", tenant_id) \
            .execute()
        print(f"  Marked {len(existing_versions)} previous version(s) as not latest")
    else:
        new_version = 1

    path = Path(file_path)
    pages = extract_text(path)

    if not pages:
        return False, "Could not extract text from this file. It may be scanned, image-based, or empty."

    if check_prompt_injection(pages):
        os.remove(path)
        return False, "This document was rejected because it contains content that is not allowed."

    if not check_relevance(pages):
        os.remove(path)
        return False, "This document does not appear to be relevant to the construction knowledge base. Only roofing, glazing, building codes, safety, and related technical documents are accepted."

    chunks = chunk_pages(pages, filename)
    chunks = embed_chunks(chunks)
    store_chunks(chunks, tenant_id, version=new_version, is_latest=True, original_filename=filename)
    register_document(filename, len(chunks), tenant_id,
                      version=new_version, is_latest=True, original_filename=filename)

    print(f"Ingested {len(chunks)} chunks from {filename} (version {new_version})")
    return True, {"chunks": len(chunks), "version": new_version}


# Backward-compatible alias
ingest_single_pdf = ingest_single_file

if __name__ == "__main__":
    ingest_all()